# app/services/job_post/upload_jd/uplod_job_post.py

import os
import re
import ast
import json
import base64
import asyncio
import hashlib
import traceback
import redis.asyncio as redis

try:
    import fitz
except Exception:
    fitz = None

from io import BytesIO
from docx import Document
from datetime import datetime
from agents import Agent, Runner, set_tracing_disabled

from typing import Dict, Union, Any
from pdf2image import convert_from_bytes
from fastapi import UploadFile, File, HTTPException, status
 
from app.config.app_config import AppConfig
from app.prompts.extract_jd_prompt import prompt_template
from app.services.job_post.upload_jd.base import BaseUploadJobPost
 
settings = AppConfig()

try:
    # Groq inference works via OpenAI-compatible API. Disable OpenAI tracing to avoid non-fatal 401 noise.
    set_tracing_disabled(True)
except Exception:
    pass
 
class UploadJobPost(BaseUploadJobPost):
    """
    Handles the extraction of job details from PDF and DOCX files using an LLM agent
    and a Redis cache to avoid redundant API calls.
    """
    def __init__(self, redis_store: redis.Redis):
        """Initializes the service with a Redis client instance."""
        self.redis_store = redis_store

    def create_agent(self) -> Union[Agent, Dict]:
        """Initializes and returns the LLM agent using Groq via OpenAI-compatible settings."""
        try:
            llm_api_key = settings.effective_groq_api_key
            if not llm_api_key:
                raise ValueError("Groq API key is not configured")

            os.environ['OPENAI_API_KEY'] = llm_api_key
            os.environ['OPENAI_BASE_URL'] = settings.groq_base_url
            return Agent(
                name="JD Extractor",
                model=settings.effective_groq_model,
                instructions=prompt_template,
            )
        except Exception as e:
            error_str = str(e)
            if "api key" in error_str.lower() or "authentication" in error_str.lower() or "unauthorized" in error_str.lower():
                return {"error": "LLM API key is invalid or unauthorized. Please check your credentials."}
            return {"error": "We couldn’t initialize the extraction model. Please try again later."}
 
    def extract_text_from_docx(self, docx_bytes: bytes) -> str:
        """Extracts text content from a DOCX file."""
        try:
            docx_stream = BytesIO(docx_bytes)
            doc = Document(docx_stream)
            text_content = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            full_text = "\n".join(text_content)
            if full_text.strip():
                return full_text

            else:
                return None

        except Exception as e:
            return None

    def extract_text_from_pdf(self, pdf_bytes: bytes) -> Union[str, None]:
        """Extracts selectable text from a PDF using PyMuPDF."""
        if fitz is None:
            print("[WARN] PyMuPDF (fitz) is not available. Skipping direct PDF text extraction.")
            return None

        try:
            with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf_doc:
                page_count = len(pdf_doc)
                if page_count == 0:
                    return None

                pages_to_process = min(page_count, settings.max_pdf_pages)
                extracted_chunks = []
                for page_index in range(pages_to_process):
                    page_text = pdf_doc[page_index].get_text("text")
                    if page_text and page_text.strip():
                        extracted_chunks.append(page_text.strip())

            if extracted_chunks:
                return "\n".join(extracted_chunks)

            return None

        except Exception as e:
            print(f"[WARN] PDF text extraction failed with PyMuPDF: {e}")
            return None

    def _extract_first_json_object(self, raw_text: str) -> Union[str, None]:
        """Extracts the first balanced JSON object from a noisy string."""
        if not raw_text:
            return None

        start_index = raw_text.find("{")
        while start_index != -1:
            depth = 0
            in_string = False
            escape = False

            for idx in range(start_index, len(raw_text)):
                ch = raw_text[idx]

                if in_string:
                    if escape:
                        escape = False
                    elif ch == "\\":
                        escape = True
                    elif ch == '"':
                        in_string = False
                    continue

                if ch == '"':
                    in_string = True
                elif ch == "{":
                    depth += 1
                elif ch == "}":
                    depth -= 1
                    if depth == 0:
                        return raw_text[start_index : idx + 1]

            start_index = raw_text.find("{", start_index + 1)

        return None

    def _parse_agent_json_output(self, raw_output: str) -> Union[Dict, None]:
        """Parses JSON from LLM output, tolerating code fences and minor formatting noise."""
        if not raw_output:
            return None

        candidates = []

        fenced_match = re.search(
            r"```(?:jsonc?|)?\s*(.*?)```",
            raw_output,
            re.DOTALL | re.IGNORECASE,
        )
        if fenced_match:
            candidates.append(fenced_match.group(1).strip())

        extracted_object = self._extract_first_json_object(raw_output)
        if extracted_object:
            candidates.append(extracted_object.strip())

        candidates.append(raw_output.strip())

        seen = set()
        for candidate in candidates:
            if not candidate or candidate in seen:
                continue
            seen.add(candidate)

            normalized_variants = [
                candidate,
                candidate.replace("“", '"').replace("”", '"').replace("’", "'").replace("‘", "'"),
            ]

            for variant in normalized_variants:
                cleaned = re.sub(r",\s*([}\]])", r"\1", variant)

                try:
                    parsed = json.loads(cleaned)
                    if isinstance(parsed, dict):
                        return parsed
                except json.JSONDecodeError:
                    pass

                literal_candidate = re.sub(r"\btrue\b", "True", cleaned, flags=re.IGNORECASE)
                literal_candidate = re.sub(r"\bfalse\b", "False", literal_candidate, flags=re.IGNORECASE)
                literal_candidate = re.sub(r"\bnull\b", "None", literal_candidate, flags=re.IGNORECASE)

                try:
                    parsed_literal = ast.literal_eval(literal_candidate)
                    if isinstance(parsed_literal, dict):
                        return parsed_literal
                except (ValueError, SyntaxError):
                    pass

        return None

    def _coerce_int(self, value: Any, default: int = 0) -> int:
        if value is None:
            return default
        if isinstance(value, (int, float)):
            return int(value)
        digits = re.sub(r"[^0-9]", "", str(value))
        if not digits:
            return default
        try:
            return int(digits)
        except ValueError:
            return default

    def _coerce_bool(self, value: Any) -> bool:
        if isinstance(value, bool):
            return value
        if isinstance(value, (int, float)):
            return bool(value)
        text = str(value).strip().lower()
        if text in {"true", "yes", "y", "1"}:
            return True
        if text in {"false", "no", "n", "0"}:
            return False
        return False

    def _normalize_skills(self, skills: Any) -> list:
        if not skills:
            return []
        if isinstance(skills, dict):
            skills = [skills]
        if not isinstance(skills, list):
            return []

        normalized = []
        for item in skills:
            if isinstance(item, str):
                skill_name = item.strip()
                if skill_name:
                    normalized.append({"skill": skill_name, "weightage": 5})
                continue
            if isinstance(item, dict):
                skill_name = (
                    item.get("skill")
                    or item.get("skill_name")
                    or item.get("name")
                    or item.get("title")
                )
                if not skill_name:
                    continue
                weight = item.get("weightage", item.get("weight", 5))
                normalized.append({
                    "skill": str(skill_name).strip(),
                    "weightage": max(1, min(10, self._coerce_int(weight, 5))),
                })
        return normalized

    def _normalize_job_details(self, job_details: Dict[str, Any]) -> Dict[str, Any]:
        if not isinstance(job_details, dict):
            return job_details

        normalized = dict(job_details)

        def pick(*keys: str) -> Any:
            for key in keys:
                if key in normalized and normalized[key] not in (None, ""):
                    return normalized[key]
            return None

        title = pick("job_title", "jobTitle", "title", "position_title")
        if title is not None:
            normalized["job_title"] = str(title).strip()

        description = pick("job_description", "jobDescription", "description")
        if description is not None:
            normalized["job_description"] = str(description).strip()

        location = pick("job_location", "jobLocation", "location", "jobLocationName")
        if location is not None:
            normalized["job_location"] = str(location).strip()

        work_from_home = pick("work_from_home", "workFromHome", "remote", "is_remote", "work_from_home_or_hybrid")
        if work_from_home is not None:
            normalized["work_from_home"] = self._coerce_bool(work_from_home)

        min_exp = pick("minimum_experience", "min_experience", "minExperience", "experience_min")
        max_exp = pick("maximum_experience", "max_experience", "maxExperience", "experience_max")
        if min_exp is not None:
            normalized["min_experience"] = self._coerce_int(min_exp, 0)
            normalized["minimum_experience"] = normalized["min_experience"]
        if max_exp is not None:
            normalized["max_experience"] = self._coerce_int(max_exp, 0)
            normalized["maximum_experience"] = normalized["max_experience"]

        skills = pick("skills_required", "skillsRequired", "skills", "skillset")
        if skills is not None:
            normalized["skills_required"] = self._normalize_skills(skills)

        return normalized

    def _has_meaningful_skills(self, skills: Any) -> bool:
        if not isinstance(skills, list):
            return False
        for item in skills:
            if not isinstance(item, dict):
                continue
            if str(item.get("skill") or "").strip():
                return True
        return False

    def _is_empty_job_details(self, job_details: Dict[str, Any]) -> bool:
        if not isinstance(job_details, dict):
            return True
        title = str(job_details.get("job_title") or "").strip()
        description = str(job_details.get("job_description") or "").strip()
        location = str(job_details.get("job_location") or "").strip()
        min_exp = self._coerce_int(job_details.get("min_experience"), 0)
        max_exp = self._coerce_int(job_details.get("max_experience"), 0)
        skills_ok = self._has_meaningful_skills(job_details.get("skills_required"))
        return not (title or description or location or skills_ok or min_exp > 0 or max_exp > 0)

    def _guess_job_title(self, text: str) -> str:
        if not text:
            return ""
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        for line in lines[:30]:
            if re.search(r"\b(job\s*title|position|role)\b", line, re.IGNORECASE):
                if ":" in line:
                    return line.split(":", 1)[-1].strip()
                return line
        return lines[0] if lines else ""

    def _guess_experience_range(self, text: str) -> tuple[int, int]:
        if not text:
            return 0, 0
        range_match = re.search(r"(\d+)\s*(?:-|to)\s*(\d+)\s*(?:\+?\s*)?(?:years|yrs)", text, re.IGNORECASE)
        if range_match:
            return self._coerce_int(range_match.group(1), 0), self._coerce_int(range_match.group(2), 0)
        single_match = re.search(r"(\d+)\s*(?:\+?\s*)?(?:years|yrs)", text, re.IGNORECASE)
        if single_match:
            years = self._coerce_int(single_match.group(1), 0)
            return years, years
        return 0, 0

    def _fallback_from_text(self, text_content: str) -> Dict[str, Any]:
        title = self._guess_job_title(text_content)
        min_exp, max_exp = self._guess_experience_range(text_content)
        work_from_home = bool(re.search(r"\b(remote|work from home|wfh|hybrid)\b", text_content, re.IGNORECASE))
        return {
            "job_title": title,
            "job_description": text_content.strip()[:8000],
            "skills_required": [],
            "job_location": "Work from home" if work_from_home else "",
            "work_from_home": work_from_home,
            "min_experience": min_exp,
            "max_experience": max_exp,
            "minimum_experience": min_exp,
            "maximum_experience": max_exp,
        }
 
    async def extract_job_details_with_text(self, text_content: str) -> Dict:
        """Processes text content with the LLM agent."""
        agent = self.create_agent()
        if isinstance(agent, dict) and 'error' in agent:
            return agent
        try:
            result = await Runner.run(agent, [{"role": "user", "content": f"Extract complete job description with skills and their weightage from the following text:\n\n{text_content}"}])
            if not result.final_output:
                return { "error": "We couldn’t extract job details from the text content. Please check the file and try again."}

            parsed_data = self._parse_agent_json_output(result.final_output)
            if isinstance(parsed_data, dict):
                return parsed_data

            print(
                "[WARN] Unable to parse JD text extraction output as JSON. "
                f"Output preview: {result.final_output[:500]}"
            )
            return {"error": "Failed to parse extracted job details."}

        except Exception as e:

            error_str = str(e)
            if "429" in error_str and "insufficient_quota" in error_str:

                return {"error": "Our extraction limit has been reached. Please try again in a few minutes."}

            if "api key" in error_str.lower() or "authentication" in error_str.lower() or "unauthorized" in error_str.lower():

                return {"error": "LLM API key is invalid or unauthorized. Please check your credentials."}

            return {"error": "Failed to extract job details."}

 
    async def extract_job_details_with_agent_image(self, file_content: bytes) -> Dict:

        """Converts a PDF to images and processes with the LLM agent."""
        print(f"[INFO] Processing PDF. Converting to images for LLM analysis.")
        try:
            configured_poppler_path = (
                getattr(settings, "poppler_path", None)
                or os.getenv("POPPLER_PATH")
                or os.getenv("poppler_path")
                or os.getenv("POPLER_PATH")
                or os.getenv("popler_path")
            )

            if configured_poppler_path:
                configured_poppler_path = str(configured_poppler_path).strip().strip('"').strip("'")
            if configured_poppler_path and configured_poppler_path.lower() in {"none", "null"}:
                configured_poppler_path = None

            if configured_poppler_path:
                try:
                    images = convert_from_bytes(
                        file_content,
                        poppler_path=configured_poppler_path,
                        fmt="jpeg",
                        dpi=200,
                    )
                except Exception as poppler_err:
                    print(
                        "[WARN] Poppler conversion with configured path failed "
                        f"('{configured_poppler_path}'): {poppler_err}. Retrying using system PATH."
                    )
                    images = convert_from_bytes(file_content, fmt="jpeg", dpi=200)
            else:
                images = convert_from_bytes(file_content, fmt="jpeg", dpi=200)
        except Exception as e:
            print(f"[ERROR] PDF to image conversion failed: {e}")
            return {"error": "We couldn’t read the uploaded PDF. Please ensure the file is not corrupted and try again."}

        total_pages = len(images)
        if total_pages > settings.max_pdf_pages:
            print(f"[WARN] PDF has {total_pages} pages which exceeds the maximum allowed {settings.max_pdf_pages} pages.")
            return {"error": "The uploaded PDF exceeds the page limit, Please upload a smaller file with minimum of 3 pages."}

        if not images:
            print(f"[WARN] No images extracted from PDF.")
            return {"error": "We couldn’t extract any readable content from your PDF. Please check the file and try again."}

        agent = self.create_agent()

        if isinstance(agent, dict) and 'error' in agent:
            return {
                "error": agent.get(
                    "error",
                    "We’re having trouble setting up the extraction service. Please try again later.",
                )
            }

        all_results = []

        for i, image in enumerate(images):

            buffer = BytesIO()

            image.save(buffer, format="JPEG", optimize=True)

            image_url = f"data:image/jpeg;base64,{base64.b64encode(buffer.getvalue()).decode()}"

            try:
                print(f"[INFO] Calling agent for PDF page {i+1}.")
                result = await Runner.run(agent, [{"role": "user", "content": [

                    {"type": "input_image", "image_url": image_url},

                    {"type": "input_text", "text": "Extract the complete job description with skills and their weightage from this image. Output ONLY a valid JSON object. Do not include any text, explanations, or markdown code fences."}

                ]}], timeout = 45)

                if not result.final_output:
                    print(f"[WARN] Agent returned empty output for PDF page {i+1}. Skipping.")
                    continue

                parsed_data = self._parse_agent_json_output(result.final_output)

                if isinstance(parsed_data, dict):
                    all_results.append(parsed_data)
                else:
                    print(f"[WARN] Agent output for page {i+1} was not valid JSON. Skipping.")
                    continue

            except asyncio.TimeoutError:
                print(f"[WARN] Page {i+1} processing timed out.")
                continue

            except Exception as e:

                error_str = str(e)
                print(f"[ERROR] Agent call for page {i+1} failed. Error: {e}")
                if "429" in error_str and "insufficient_quota" in error_str:
                    return {"error": "Our extraction limit has been reached. Please try again in a few minutes."}


                if "api key" in error_str.lower() or "authentication" in error_str.lower() or "unauthorized" in error_str.lower():
                    return {"error": "There was an issue connecting to the extraction service. Please try again shortly."}

                # Continue processing remaining pages instead of failing the entire file on one page error.
                continue


        if not all_results:
            print(f"[WARN] No valid data extracted from any PDF page.")
            return {"error": "We couldn’t extract job details from the uploaded document. Please check the content and try again."}

        merged_result = {}

        for page_json in all_results:

            if isinstance(page_json, dict):

                for key, value in page_json.items():

                    if key not in merged_result:

                        merged_result[key] = value

                    elif isinstance(value, list) and isinstance(merged_result[key], list):

                        merged_result[key].extend(value)

                    elif isinstance(value, str) and isinstance(merged_result[key], str):

                        merged_result[key] += " " + value

        print(f"[INFO] Merged results from all PDF pages successfully.")
        return merged_result
 
    async def job_details_file_upload(self, file: UploadFile = File(...)) -> Dict:
        """
        The main function to handle job description file uploads.
        It checks the cache before performing expensive extraction.
        
        NOTE: job_id is intentionally NOT accepted here as it's not used in caching or extraction.
        """
        print(f"[INFO] Starting JD upload and extraction.")

       
        try:
            file_content = await file.read()
        except Exception:
            return {"error": "Please upload a proper job description file (PDF or DOCX)"}
        
        if len(file_content) == 0:
            return {"error": "Empty file uploaded. Please upload a valid job description file."}

        file_hash = hashlib.sha256(file_content).hexdigest()
        print(f"[INFO] Generated hash for file: {file_hash}")
        
        # 2. Check the cache for the hash
        try:
            cached_data = await self.redis_store.hgetall(file_hash)
            if cached_data:
                print(f"[INFO] Cache hit for hash '{file_hash}'.")
                cached_content = cached_data.get("extracted_content") or cached_data.get(b"extracted_content")
                
                if cached_content is not None:
                    # Safely decode if it's bytes, otherwise assume it's a string
                    if isinstance(cached_content, bytes):
                         cached_content = cached_content.decode('utf-8')
                         
                    print(f"[INFO] Returning cached job details.")
                    cached_details = json.loads(cached_content)
                    normalized = self._normalize_job_details(cached_details)
                    if not self._is_empty_job_details(normalized):
                        return {"job_details": normalized}
                    print(f"[WARN] Cached extraction was empty. Reprocessing file.")
                else:
                    print(f"[WARN] Cached data for hash '{file_hash}' is missing 'extracted_content'. Re-processing.")
        
        except Exception as e:
            print(f"[ERROR] Redis cache lookup failed (Critical exception during retrieval): {e}")
            # The function will proceed to the LLM extraction path.

        print(f"[INFO] Cache miss for hash '{file_hash}'. Proceeding with LLM extraction.")
        
        file_type = (file.content_type or "").lower()
        file_name = (file.filename or "").lower()

        # Some clients send octet-stream or empty content-type for regular uploads.
        if file_type in {"", "application/octet-stream"}:
            if file_name.endswith(".pdf"):
                file_type = "application/pdf"
            elif file_name.endswith(".docx"):
                file_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        job_details = None
        extracted_text = None
 
        if file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':

            if(len(file_content) > settings.max_file_size_docx):
                return {"error": "The uploaded DOCX file exceeds the maximum allowed size of 5 MB. Please upload a smaller file."}

            text_content = self.extract_text_from_docx(file_content)

            if not text_content:

                return {"error": "We couldn’t extract details from the uploaded file. Please make sure it’s a clear and readable job description (PDF or DOCX) and try again."}

            extracted_text = text_content
            job_details = await self.extract_job_details_with_text(text_content)

        elif file_type == 'application/pdf':

            if(len(file_content) > settings.max_file_size_pdf):

                return {"error": "The uploaded PDF file exceeds the maximum allowed size of 10 MB. Please upload a smaller file."}

            text_content = self.extract_text_from_pdf(file_content)
            if text_content:
                print("[INFO] Extracted selectable text from PDF. Using text-based JD extraction path.")
                extracted_text = text_content
                job_details = await self.extract_job_details_with_text(text_content)
            else:
                print("[INFO] No selectable text found in PDF. Falling back to image-based extraction path.")
                job_details = await self.extract_job_details_with_agent_image(file_content)

        else:
            print(f"[ERROR] Unsupported file type: {file_type}.")
            return {"error": "Unsupported file type. Please upload your job description as a PDF or DOCX file."}
 
        if isinstance(job_details, dict) and "error" in job_details:
            return {
                "error": job_details.get(
                    "error",
                    "We couldn’t process the job description right now. Please upload a clear and readable JD file (PDF or DOCX) and try again later.",
                )
            }
 
        if not job_details:
            print("[ERROR] Final job details extraction failed. Check the file content.")
            return {"error": "We couldn’t process the job description right now. Please try again later."}

        if isinstance(job_details, dict):
            job_details = self._normalize_job_details(job_details)
            if self._is_empty_job_details(job_details) and extracted_text:
                print("[WARN] LLM extraction returned empty fields. Falling back to text-based extraction.")
                job_details = self._fallback_from_text(extracted_text)
 
        if isinstance(job_details, dict) and job_details.get("job_description"):
            job_details['job_description'] = job_details['job_description'].replace('\n', ' ')
        
        # 3. Store the successful result in the cache
        try:
            hash_data = {
                "file_name": file.filename,
                "extracted_content": json.dumps(job_details),
                "timestamp": str(datetime.now())
            }
            await self.redis_store.hset(file_hash, mapping=hash_data)
            print(f"[INFO] Stored new job details in cache for hash: {file_hash}")
        except Exception as e:
            print(f"[ERROR] Failed to store data in Redis cache: {e}")
            pass # Failed to store in Redis, but we still return the successful result
        
        print("[INFO] Job details extracted successfully.")
        # Return a consistent dictionary format for success
        return {"job_details": job_details}